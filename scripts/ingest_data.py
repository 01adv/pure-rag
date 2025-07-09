import argparse
import pandas as pd
from docx import Document
from loguru import logger
from langchain.text_splitter import RecursiveCharacterTextSplitter
import sys
from pathlib import Path

# Add the project root to the Python path to allow importing from 'app'
sys.path.append(str(Path(__file__).parent.parent))

from app.dependencies import chroma_client
from app.services.rag import add_document


def extract_doc_content(doc_path: str) -> str:
    """Extracts all text content from a .docx file."""
    try:
        doc = Document(doc_path)
        full_text = []
        for para in doc.paragraphs:
            if text := para.text.strip():
                full_text.append(text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell_text := cell.text.strip():
                        full_text.append(cell_text)
        logger.info(f"Successfully extracted text from {doc_path}")
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Failed to extract content from {doc_path}: {e}")
        raise


def ingest_product_catalogue(collection_name: str, file_path: str):
    """Ingests product data from a CSV or Excel file into a ChromaDB collection."""
    logger.info(f"Starting ingestion for product catalogue: {file_path}")
    try:
        file_extension = Path(file_path).suffix.lower()
        if file_extension == '.csv':
            df = pd.read_csv(file_path)
        elif file_extension in ['.xlsx', '.xls']:
            df = pd.read_excel(file_path)
        else:
            raise ValueError(f"Unsupported file format for product catalogue: {file_extension}")
        
        logger.info(f"Loaded {len(df)} records from {file_path}")

        collection = chroma_client.get_or_create_collection(collection_name)
        logger.info(f"Using collection '{collection_name}'.")

        for index, row in df.iterrows():
            text = row.get("description", "")
            
            # Clean and convert price
            price_str = str(row.get("price", "0")).replace("$", "").strip()
            try:
                price_float = float(price_str)
            except (ValueError, TypeError):
                price_float = 0.0

            metadata = {
                "id": f"prod_{index}",
                "product_name": row.get("name"),
                "price": price_float,
                "benefits": row.get("benefits", ""),
                "top_ingredients": row.get("ingredients", ""),
                "benefits_of_ingredients": row.get("benefits_of_ingredients", ""),
                "reviews": row.get("reviews", ""),
            }
            add_document(collection, text, metadata)
        logger.success(f"Successfully ingested {len(df)} products into '{collection_name}'.")
    except Exception as e:
        logger.error(f"Product catalogue ingestion failed: {e}", exc_info=True)
        raise


def ingest_additional_info(collection_name: str, file_path: str):
    """Ingests and chunks additional information from a DOCX file."""
    logger.info(f"Starting ingestion for additional info: {file_path}")
    try:
        full_text = extract_doc_content(file_path)
        splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=80)
        chunks = splitter.split_text(full_text)
        logger.info(f"Split document into {len(chunks)} chunks.")

        collection = chroma_client.get_or_create_collection(collection_name)
        logger.info(f"Using collection '{collection_name}'.")

        for idx, chunk in enumerate(chunks):
            metadata = {"id": f"info_{idx}", "source": "pure.docx"}
            add_document(collection, chunk, metadata)
        logger.success(f"Successfully ingested {len(chunks)} chunks into '{collection_name}'.")
    except Exception as e:
        logger.error(f"Additional info ingestion failed: {e}", exc_info=True)
        raise


# def main():
#     """
#     Main function to run the data ingestion script from the command line.
#     """
#     parser = argparse.ArgumentParser(description="Data ingestion script for Pure minimalist.")
#     parser.add_argument(
#         "--force",
#         action="store_true",
#         help="Force re-ingestion by deleting existing ChromaDB collections.",
#     )
#     args = parser.parse_args()

#     DATA_DIR = Path(__file__).parent.parent / "data"
#     PRODUCT_CATALOGUE_PATH = DATA_DIR / "products2.csv"
#     ADDITIONAL_INFO_PATH = DATA_DIR / "pure.docx"

#     if args.force:
#         logger.warning("Force flag set. Deleting existing collections.")
#         try:
#             chroma_client.delete_collection("skincare")
#             logger.info("Deleted 'skincare' collection.")
#         except Exception:
#             logger.warning("Could not delete 'skincare' collection (it may not exist).")
#         try:
#             chroma_client.delete_collection("skincare_combined")
#             logger.info("Deleted 'skincare_combined' collection.")
#         except Exception:
#             logger.warning("Could not delete 'skincare_combined' collection (it may not exist).")

#     try:
#         ingest_product_catalogue("skincare", str(PRODUCT_CATALOGUE_PATH))
#         ingest_additional_info("skincare_combined", str(ADDITIONAL_INFO_PATH))
#         logger.success("Data ingestion completed successfully.")
#     except Exception as e:
#         logger.error(f"A critical error occurred during data ingestion: {e}")
#         sys.exit(1)


# if __name__ == "__main__":
#     main()


def main(force=False):
    DATA_DIR = Path(__file__).parent.parent / "data"
    PRODUCT_CATALOGUE_PATH = DATA_DIR / "products2.csv"
    ADDITIONAL_INFO_PATH = DATA_DIR / "pure.docx"

    if force:
        try:
            chroma_client.delete_collection("skincare")
            logger.info("Deleted 'skincare' collection.")
        except Exception:
            logger.warning("Could not delete 'skincare' collection.")
        try:
            chroma_client.delete_collection("skincare_combined")
            logger.info("Deleted 'skincare_combined' collection.")
        except Exception:
            logger.warning("Could not delete 'skincare_combined' collection.")

    try:
        ingest_product_catalogue("skincare", str(PRODUCT_CATALOGUE_PATH))
        ingest_additional_info("skincare_combined", str(ADDITIONAL_INFO_PATH))
        logger.success("Data ingestion completed successfully.")
    except Exception as e:
        logger.error(f"A critical error occurred during data ingestion: {e}")
        raise

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Data ingestion script for Pure minimalist.")
    parser.add_argument("--force", action="store_true", help="Force re-ingestion by deleting existing ChromaDB collections.")
    args = parser.parse_args()
    main(args.force)